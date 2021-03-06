import cherrypy
from cherrypy.lib.static import serve_file
import os
from io import StringIO
import sys
import dill as pickle
#import pickle
from subprocess import call

try:
    from sympy import latex
except:
    print('No sympy installed')

def section(text):
    print("<h1>"+text+"</h1>")

#Globals = {'section':section}
#Locals = {}

def getAllInside(first,last,content):
    # returns a dict of values first+value+last and the values
    valuesDict = {}
    while first in content and last in content:
        f = content.index(first)
        l = content[f:].index(last)
        key = content[f:f+l+len(last)]
        value = key.replace(first,'').replace(last,'')
        valuesDict[key] = value
        content = content.replace(key,'')
    return valuesDict

def startWith(text,content):
    return content.replace('\n','').replace(' ','')[:len(text)] == text

def emptyLine(content):
    if content.replace('\n','').replace(' ','') == '':
        return True
    else:
        return False

class Capturing(list):
    def __enter__(self):
        self._stdout = sys.stdout
        sys.stdout = self._stringio = StringIO()
        return self
    def __exit__(self, *args):
        self.extend(self._stringio.getvalue().splitlines())
        print(self)
        sys.stdout = self._stdout
	  
class WillNotebook(object):
    emptyLineSymbol = '.'
    def __init__(self,**kwargs):
        super(WillNotebook).__init__(**kwargs)
        self.Globals = {'section':section}
        self.Locals = {}

    @cherrypy.expose
    def index(self):
        self.Globals = {'section':section}
        self.Locals = {}
        notebook = open('notebook.html','rb')
        self.archive = {'page':[]}
        return notebook

    @cherrypy.expose
    def newCell(self,index):
        print('New cell called')
        self.archive['page'].insert(int(index),{'content':'','output':'.'})
        return 'Cell inserted'

    @cherrypy.expose
    def deleteCell(self,index):
        index = int(index)
        print('Delete cell called')
        print('Index: ',index)
        if 'type' in self.archive['page'][index]['content']:
            if self.archive['page'][index]['content']['type'] == 'image':
                filename = self.archive['page'][index]['content']['img']
                os.remove(os.getcwd()+'/Archieves/Images/'+filename)
        self.archive['page'].pop(int(index))
        return 'Cell deleted'

    @cherrypy.expose
    def evalCell(self,cell,content):
        cell = int(cell)
        print(cell)
        if cell == len(self.archive['page']):
            self.archive['page'].append({'content':content,'output':'.'})
        self.archive['page'][cell]['content'] = content
        if startWith('#code',content):
            output = self.handlePythonCode(content)
            ## {{}} soh nao sera avaliado em #code
        else:
            if '{{' in content and '}}' in content:
                print('TEMMM')
                content = self.handleValues(content)
            if startWith('!#',content):
                output = self.handleSections(content)
            elif startWith('!eq',content):
                output = self.handleEquations(content)
            elif startWith('!title',content):
                output = self.handleTitle(content)
            else:
                output = content
        if emptyLine(output):
            output = self.emptyLineSymbol
        self.archive['page'][cell]['output'] = output
        print('Out: ',output)
        return output

    def handlePythonCode(self,content):
        with Capturing() as output:
            try:
                exec(content,self.Globals,self.Locals)
            except Exception as e:
                print('Exception: ', e)
        if output == []:
            label = ''
            for line in content.split('\n'):
                if '#code' in line:
                    label = line.replace('#code','')
                    break
            return '<font class="dontprint" color="green">[code]'+label+'</font>'
        else:
            # Retorna o primeiro indice da lista (str)
            return output[0]

    def handleValues(self,content):
        toEvaluate = getAllInside('{{','}}',content)
        for expr in toEvaluate:
            try:
                out = eval(toEvaluate[expr],self.Globals,self.Locals)
                objType = str(type(out))
                if 'sympy' in objType or 'list' in objType and 'sympy' in str(type(out[0])):
                    print('Is a sympy object!')
                    try:
                        out = str(latex(out))
                        if not '!eq' in content:
                            out = '$'+out+'$'
                    except:
                        print('Latex error')
                        out = str(out)
                else: 
                    print('Not a sympy object! It it: ',objType)
                    out = str(out)
                content = content.replace(expr,out)
            except Exception as e:
                print('excep: ',e)
                content = str(e)
        return content

    def handleSections(self,content):
        n = 1 #number of #'s in the section mark
        pos = content.index('!#')+2
        for char in content[pos:]:
            if char == '#':
                n+=1
            else:
                break
        heading = content[pos+n:]
        return '<h'+str(n)+'>'+heading+'</h'+str(n)+'>'

    def handleEquations(self,content):
        pos = content.index('!eq')+3
        eqContent = ''
        for line in content.split('\n'):
            if '!eq' in line:
                label = line.replace('!eq ','')
            else:
                eqContent += line
        eq = '\\begin{equation}'
        if label.replace('<label>',''):
            eq += '\label{'+label+'}\n'
        eq += eqContent+'\end{equation}'
        return eq

    def handleTitle(self,content):
        title = content.replace('!title ','')
        return '<center><b><font size="7">'+title+'</font></b></center><br><br>'
    
    @cherrypy.expose
    def image(self,cell,img,label,source,caption):
        cell = int(cell)
        if cell == len(self.archive['page']):
            filename = img.filename
            self.archive['page'].append({'content':{'type':'image','img':filename,'label':label,'source':source,'caption':caption},'output':'.'})
        i = open(os.getcwd()+'/Archieves/Images/'+filename,'wb')
        while True:
            data = img.file.read(4096)
            if not data:
                break
            else:
                i.write(data)
            print('Loading...')
        i.close()
        output = '<br><center><figcaption>'+caption+'</figcaption>'+'<img style="max-width:800px" src="Archieves/Images/'+filename+'"><br>'+source+'</center><br>'
        self.archive['page'][cell]['output'] = output
        return output

    @cherrypy.expose
    def saveFile(self,filename,extension):
        if extension == 'will':
            self.saveAsWill(filename)
        elif extension == 'tex':
            self.saveAsTex(filename)
        elif extension == 'docx':
            self.saveAsDocx(filename)
        elif extension == 'pdflatex':
            self.saveAsPdfLatex(filename)
            extension = 'pdf'
        return serve_file(os.getcwd()+'/Archieves/'+filename+'.'+extension,"application/x-download","attachment")

    def saveAsWill(self,filename):
        archive = open(os.getcwd()+'/Archieves/'+filename+'.will','wb')
        try:
            self.archive['Locals'] = self.Locals
        except Exception as e:
            print('Could not save state! ',e)
        pickle.dump(self.archive,archive)
        archive.close()

    def saveAsTex(self,filename,article=True):
        archive = open(os.getcwd()+'/Archieves/'+filename+'.tex','w')
        if article:
            texClass = 'article'
        else:
            texClass = 'report'
        archive.write('''\\documentclass{'''+texClass+'''}
\\usepackage[T1]{fontenc}
\\usepackage[utf8]{inputenc}
\\usepackage{graphicx}
\\usepackage{mathtools}

\\setcounter{secnumdepth}{5}
\\setcounter{tocdepth}{5}

\\begin{document}

''')
        for cell in self.archive['page']:
            content = cell['content']
            show = True
            if '!# ' in content:
                if article:
                    content = '\section{'
                else:
                    content = '\chapter{'
                content += cell['content'].replace('!# ','')
                content += '}'
            elif '!## ' in content:
                if article:
                    content = '\subsection{'
                else:
                    content = '\section{'
                content += cell['content'].replace('!## ','')
                content += '}'
            elif '!### ' in content:
                if article:
                    content = '\subsubsection{'
                else:
                    content = '\subsection{'
                content += cell['content'].replace('!### ','')
                content += '}'
            elif '!#### ' in content:
                if article:
                    content = '\paragraph{'
                else:
                    content = '\subsubsection{'
                content += cell['content'].replace('!#### ','')
                content += '}'
            elif '!##### ' in content:
                if article:
                    content = '\subparagraph{'
                else:
                    content = '\paragraph{'
                content += cell['content'].replace('!##### ','')
                content += '}'
            elif '[code]' in cell['output']:
                show = False
            else:
                content = cell['output']
            if show:
                archive.write(content+'\n\n')
        archive.write('\\end{document}')
        archive.close()

    def saveAsPdfLatex(self,filename):
        self.saveAsTex(filename)
        call(['pdflatex','-interaction=nonstopmode',filename+'.tex'], cwd=os.getcwd()+'/Archieves/')
        call(['bibtex','-interaction=nonstopmode',filename+'.aux'], cwd=os.getcwd()+'/Archieves/')
        call(['pdflatex','-interaction=nonstopmode',filename+'.tex'], cwd=os.getcwd()+'/Archieves/')
        call(['pdflatex','-interaction=nonstopmode',filename+'.tex'], cwd=os.getcwd()+'/Archieves/')

    def saveAsDocx(self,filename):
        from docx import Document
        from docx.shared import Inches
        d = Document()
        for cell in self.archive['page']:
            content = cell['content']
            if '!# ' in content:
                d.add_heading(content.replace('!# ',''),level=1)
            elif '!## ' in content:
                d.add_heading(content.replace('!## ',''),level=2)
            elif '!### ' in content:
                d.add_heading(content.replace('!### ',''),level=3)
            elif '!#### ' in content:
                d.add_heading(content.replace('!#### ',''),level=4)
            elif '!##### ' in content:
                d.add_heading(content.replace('!##### ',''),level=5)
            else:
                d.add_paragraph(content)

        d.save(os.getcwd()+'/Archieves/'+filename+'.docx')

    @cherrypy.expose
    def openFile(self,filename):
        self.archive = {}
        print('Openning file ',filename)
        try:
            archive = open(os.getcwd()+'/Archieves/'+filename,'rb')
        except:
            content = 'Error loading file. File not found.'
            output = content
            print('Error opening file. File does not exist!')
            return '<center id="c1"><textarea id="1" action="evalCell" style="width: 800px; display: none;">'+content+'</textarea></center><center id="co1"><div id="o1" style="width: 800px; text-align: justify;">'+output+'</div></center>'
        self.archive = pickle.load(archive)
        try:
            self.Locals = self.archive['Locals']
        except Exception as e:
            print('File has no Locals. ',e)
        archive.close()
        notebook = ''
        for cell,stuff in enumerate(self.archive['page']):
            # toDo check if cell is 'image'. If so, construct the image, not
            # the textArea
            print(cell)
            if 'type' in stuff['content']:
                if stuff['content']['type'] == 'image':
                    img = stuff['content']['img']
                    label = stuff['content']['label']
                    source = stuff['content']['source']
                    caption = stuff['content']['caption']

                    notebook += '<center id="c'+str(cell)+'"><form id="F'+str(cell)+'" enctype="multipart/form-data" method="POST" action="image" style="display: none;"><input type="file" name="img" value="Images/'+img+'" id="'+str(cell)+'" style="display: none;"><br>Label:<input name="label" value="'+label+'" id="L'+str(cell)+'"><br>Caption:<input name="caption" value="'+caption+'" id="C'+str(cell)+'"><br>Source:<input name="source" value="'+source+'" id="S'+str(cell)+'"></form></center>'
                    output = stuff['output']
                    notebook += '<center tabindex="0" id="co'+str(cell)+'"><div id="o'+str(cell)+'" style="width: 800px; text-align: justify;">'+output+'</div></center>'
            else:
                content = stuff['content']
                output = stuff['output']
                notebook += '<center id="c'+str(cell)+'"><textarea id="'+str(cell)+'" action="evalCell" style="width: 800px; display: none;">'+content+'</textarea></center>'
                notebook += '<center tabindex="0" id="co'+str(cell)+'"><div id="o'+str(cell)+'" style="width: 800px; text-align: justify;">'+output+'</div></center>'
        return notebook


if __name__ == '__main__':
    conf = {'/':{'tools.sessions.on':True,'tools.staticdir.on':True,'tools.staticdir.dir':os.getcwd()}}
    cherrypy.quickstart(WillNotebook(),'/',config=conf)
